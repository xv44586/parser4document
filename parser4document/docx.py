# -*- coding: utf-8 -*-
# @Date    : 2020/9/21
# @Author  : mingming.xu
# @Email   : mingming.xu@zhaopin.com
# @File    : docx.py
import os
from PIL import Image
from io import BytesIO
import docx
import pandas as pd

# `docx.Document` to open or create a document
from docx.document import Document
# w:p, w:tb1
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
# Table : Proxy class for a WordprocessingML ``<w:tbl>`` element.
# _Cell : Table cell
from docx.table import _Cell, Table
# Paragraph : Proxy object wrapping ``<w:p>`` element.
from docx.text.paragraph import Paragraph

from .utils import format


class DocXParser(object):
    """
    parse document content
    """

    def __init__(self, word_path):
        self.word_path = word_path
        self.document = docx.Document(word_path)

    def parse_paragraph(self, paragraph):
        """
        paragraph对象是document中的主要对象，其中存储了内容与对应样式；
        paragraph又可以细分为多个run对象，每个run对象内又包含对应的内容与样式

        ref: [paragraph-objects](https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects)
        """
        runs = paragraph.runs
        parse_runs = []
        if len(runs) < 1:
            return []

        for run in runs:
            parse_run = {
                'text': format(run.text),
                'bold': run.bold,
                'underline': run.underline,
                'italic': run.italic,
                'highlight_color': run.font.highlight_color,
                'strike': run.font.strike,
                'double_strike': run.font.double_strike,
                'outline': run.font.outline,
                'shadow': run.font.shadow,
                'font_name': run.font.name if run.font else None,
                'font_size': run.font.size.pt if run.font.size else None
            }
            if len(parse_runs) > 0 and self.compare_style(parse_runs[-1], parse_run):
                parse_runs[-1] = self.combine_runs(parse_runs[-1], parse_run)
            else:
                parse_runs.append(parse_run)

        return parse_runs

    @staticmethod
    def compare_style(run_1, run_2):
        is_same = True
        for type_name, value in run_1.items():
            if type_name == 'text':
                continue
            if value != run_2[type_name]:
                is_same = False
                break
        return is_same

    @staticmethod
    def combine_runs(run_1, run_2):
        run_1['text'] = run_1['text'] + run_2['text']
        return run_1

    def parse_cell(self, cell):
        """
        cell 是table对象的主要内容，cell内包含对应的内容和一个paragraph list

        ref: [table](https://python-docx.readthedocs.io/en/latest/api/table.html#id1)
        todo: 合并样式相同的相邻paragraph
        """
        parse_para = []
        for paragraph in cell.paragraphs:
            parse_runs = self.parse_paragraph(paragraph)
            parse_para.append({'runs': parse_runs})

        return {'text': cell.text, 'paragraphs': parse_para}

    def read_table(self, table):
        """
        解析table对象中的内容
        :param table:
        :return:
        """
        table = [[self.parse_cell(cell) for cell in row.cells] for row in table.rows]
        cell_len = len(table[0])
        row_len = len(table)

        table = list(reversed([list(reversed(row)) for row in table]))
        new_table = []
        for row_idx, row in enumerate(table):
            new_row = []
            for cell_idx, cell in enumerate(row):
                # 与上一行内容相同
                #             if row_idx < row_len - 1 and cell['text'] == table[row_idx + 1][cell_idx]['text']:
                #                 cell['text'] = ' '
                #                 new_row.append(cell)
                #             else:

                #                 new_row.append(cell)
                # 与上一列内容相同
                if cell_idx < cell_len - 1 and cell['text'] == row[cell_idx + 1]['text']:
                    cell['text'] = ' '
                    new_row.append(cell)
                else:
                    new_row.append(cell)
                # 与上一行与上一列都相同，保留原始内容
            #             if row_idx < row_len -1 and cell_idx < cell_len -1 and cell == table[row_idx + 1][cell_idx] and cell == row[cell_idx + 1]:
            #                 new_row[cell_idx] = cell

            new_table.append(new_row)

        new_table = list(reversed([list(reversed(row)) for row in new_table]))

        return new_table

    def iter_block_items(self, parent):
        """切分document"""
        if isinstance(parent, Document):  # 是doc
            parent_elm = parent.element.body  # 返回文档的内容 w:body
        elif isinstance(parent, _Cell):  # 是表格单元就将表格单元的内容返回
            parent_elm = parent._tc  # w:tc [table cell]
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():  # 迭代子元素，【分割块】
            if isinstance(child, CT_P):  # 属于w:p
                yield Paragraph(child, parent)  # 生成器生成段落代理【标记这个块是段落】
            elif isinstance(child, CT_Tbl):  # 属于w:tb1
                yield Table(child, parent)  # 生成器生成表格代理【标记这个块是表格】

    def read(self, with_textbox=False, verbose=False):
        """
        textbox 不在基本API下，而是通过直接从xml里寻找tag 抽取，所以需要额外指定
        """
        if not hasattr(self, 'document'):
            self.document = docx.Document(self.word_path)
        blocks = []
        for block in self.iter_block_items(self.document):
            # 对应的打印输出
            if isinstance(block, Paragraph):
                runs = self.parse_paragraph(block)
                content = format(block.text)
                blocks.append({'type': 'text', 'content': content, 'runs': runs})
                if verbose:
                    print(content)  # 是段落

            elif isinstance(block, Table):
                table_text = self.read_table(block)
                blocks.append({'type': 'table', 'content': table_text})
                text = '\n\n'.join(['\t'.join([c['text'] for c in l]) for l in table_text])
                if verbose:
                    print(text)  # 是表格

        # extract textbox
        if with_textbox:
            textbox = self.extract_textbox()
            blocks.append({'type': 'textbox', 'content': textbox, 'runs': []})

        return blocks

    @staticmethod
    def show_bytes(bytes_blobs):
        for bytes_blob in bytes_blobs:
            stream = BytesIO(bytes_blob)
            Image.open(stream)

    def extract_pics(self, show=False):
        """docx中的附件内容存储在对应的rel中，根据名称可以判断其类型，具体内容查看：
        ```python
        print(Document(file_path).part._rels.items())
        ```
        """
        blobs = []
        for _, rel in self.document.part._rels.items():
            if 'media' in rel.target_ref:
                blob = rel.target_part.blob
                blobs.append(blob)
                print(rel.target_ref)

        if show:
            self.show_bytes(blobs)
        return blobs

    def extract_textbox(self):
        """
        直接从xml里的tag里抽取textbox 中内容
        """
        children = self.document.element.body.iter()
        child_iters = []
        tags = []
        for child in children:
            # 通过类型判断目录
            if child.tag.endswith(('AlternateContent', 'textbox')):
                for ci in child.iter():

                    tags.append(ci.tag)
                    if ci.tag.endswith(('main}r', 'main}pPr')):
                        child_iters.append(ci)
        text = ['']
        for ci in child_iters:
            if ci.tag.endswith('main}pPr'):
                text.append('')
            else:
                text[-1] += ci.text
            ci.text = ''

        if len(text) < 2 and text[0] == '':
            return None

        return format(text)
